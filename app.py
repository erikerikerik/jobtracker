"""
JobTracker - Intelligent Job Search & Tracking Application
Run: python app.py
Then open: http://localhost:5000
"""

import os
import json
import sqlite3
import re
from datetime import datetime
from flask import Flask, request, jsonify, send_from_directory
import requests
from bs4 import BeautifulSoup

app = Flask(__name__, static_folder='static', template_folder='templates')

DB_PATH = os.path.join(os.path.dirname(__file__), 'data', 'jobs.db')
RESUME_PATH = os.path.join(os.path.dirname(__file__), 'data', 'resume.txt')

# ──────────────────────────────────────────────
# Database setup
# ──────────────────────────────────────────────

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conn = get_db()
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS jobs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            company TEXT NOT NULL,
            location TEXT,
            remote INTEGER DEFAULT 0,
            salary TEXT,
            url TEXT,
            source TEXT DEFAULT 'manual',
            status TEXT DEFAULT 'interested',
            score REAL,
            match_reasons TEXT,
            notes TEXT,
            date_added TEXT,
            date_applied TEXT,
            date_updated TEXT,
            raw_description TEXT
        );

        CREATE TABLE IF NOT EXISTS sources (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            url TEXT NOT NULL,
            type TEXT DEFAULT 'greenhouse',
            active INTEGER DEFAULT 1,
            last_fetched TEXT,
            job_count INTEGER DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS resume_data (
            id INTEGER PRIMARY KEY,
            content TEXT,
            skills TEXT,
            last_updated TEXT
        );
    """)
    conn.commit()

    # Seed a handful of example sources — add your own via the UI
    default_sources = [
            # Greenhouse examples
            ("Stripe",      "https://boards.greenhouse.io/stripe",      "greenhouse", 1),
            ("Brex",        "https://boards.greenhouse.io/brex",        "greenhouse", 1),
            ("Ramp",        "https://boards.greenhouse.io/ramp",        "greenhouse", 1),
            ("Plaid",       "https://boards.greenhouse.io/plaid",       "greenhouse", 1),
            ("Airbnb",      "https://boards.greenhouse.io/airbnb",      "greenhouse", 1),
            # Lever examples
            ("Mercury",     "https://jobs.lever.co/mercury",            "lever",      1),
            # Aggregators
            ("RemoteOK",    "https://remoteok.com/api",                 "remoteok",   1),
        ]
    for name, url, type_, active in default_sources:
        exists = conn.execute("SELECT id FROM sources WHERE name=?", (name,)).fetchone()
        if not exists:
            conn.execute(
                "INSERT INTO sources (name, url, type, active) VALUES (?,?,?,?)",
                (name, url, type_, active)
            )
    conn.commit()

    conn.close()

# ──────────────────────────────────────────────
# Resume handling
# ──────────────────────────────────────────────

def load_resume_text():
    conn = get_db()
    row = conn.execute("SELECT content, skills FROM resume_data WHERE id=1").fetchone()
    conn.close()
    if row:
        return row['content'], json.loads(row['skills'] or '[]')
    return "", []

# ── Generic text helpers ──────────────────────────────────────────────────────

_STOPWORDS = {
    'a', 'an', 'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
    'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through', 'during',
    'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had',
    'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might',
    'must', 'can', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she',
    'we', 'they', 'it', 'as', 'if', 'than', 'so', 'yet', 'both', 'either',
    'not', 'no', 'nor', 'too', 'very', 'just', 'our', 'your', 'their', 'its',
    'my', 'his', 'her', 'who', 'which', 'what', 'when', 'where', 'how', 'all',
    'any', 'each', 'every', 'more', 'most', 'other', 'some', 'such', 'only',
    'also', 'well', 'new', 'good', 'high', 'including', 'across', 'within',
    'experience', 'skills', 'ability', 'strong', 'excellent', 'years', 'year',
    'team', 'role', 'responsibilities', 'qualifications', 'required', 'preferred',
    'minimum', 'position', 'candidate', 'company', 'business', 'using',
    'ensure', 'support', 'provide', 'develop', 'manage', 'working', 'related',
    'relevant', 'job', 'opportunity', 'looking', 'join', 'based',
}

def _tokenize(text):
    """Tokenize text into meaningful lowercase words, stripping stopwords."""
    words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9+#\-\.]*\b', text.lower())
    return [w for w in words if len(w) > 2 and w not in _STOPWORDS]


def extract_skills_from_resume(text):
    """
    Extract key skills and terms from resume text for job matching.
    Works for any industry — pulls abbreviations, recurring terms, and
    capitalized technical/proper nouns from the resume content.
    """
    from collections import Counter

    # ALL-CAPS abbreviations (credentials, tools, acronyms — e.g. CPA, AWS, SQL, MBA)
    abbrevs = re.findall(r'\b[A-Z]{2,6}\b', text)

    # Recurring significant tokens — words that appear ≥2× are likely important
    tokens = _tokenize(text)
    counts = Counter(tokens)
    recurring = {w for w, c in counts.items() if c >= 2 and len(w) > 3}

    skills = set(a.lower() for a in abbrevs if len(a) >= 2)
    skills.update(recurring)

    # Drop overly generic words that sneak through
    noise = {'work', 'time', 'also', 'make', 'used', 'use', 'like', 'one',
             'two', 'three', 'four', 'five', 'including', 'within', 'across',
             'other', 'they', 'that', 'this', 'with', 'from', 'have', 'been'}
    skills -= noise

    return sorted(list(skills))[:100]

def estimate_competition(job_title, company):
    """Estimate applicant competition level based on company brand recognition and seniority."""
    title_lower = job_title.lower()
    company_lower = company.lower()

    # Tier-1 household-name tech/consumer brands attract the most applicants
    high_brand = [
        'google', 'apple', 'meta', 'amazon', 'microsoft', 'netflix', 'spotify',
        'stripe', 'airbnb', 'uber', 'lyft', 'coinbase', 'robinhood', 'openai',
        'slack', 'shopify', 'twitter', 'instagram', 'linkedin', 'salesforce',
    ]
    # Consulting/professional-services firms — moderately high volume
    consulting = [
        'mckinsey', 'deloitte', 'kpmg', 'pwc', 'accenture', 'bain',
        'boston consulting', 'ernst', 'guidehouse', 'treliant',
    ]

    if any(b in company_lower for b in high_brand):
        return '🔴 HIGH (~500–1,000+ applicants)'
    if any(b in company_lower for b in consulting):
        return '🟡 MEDIUM (~200–400 applicants)'
    # Senior/executive roles attract fewer but more targeted applicants
    if any(k in title_lower for k in ['vp ', 'vice president', 'director', 'head of', 'principal', 'avp']):
        return '🟡 MEDIUM (~200–350 applicants)'
    # Smaller / less well-known companies
    if len(company) < 12:
        return '🟢 MEDIUM-LOW (~100–250 applicants)'
    return '🟡 MEDIUM (~150–300 applicants)'


def score_job_against_resume(job_title, job_desc, resume_text, resume_skills):
    """
    Score a job against the resume using content overlap.
    Returns (score 0-100, list of reasons).

    Scoring is domain-agnostic — it measures how much vocabulary the job
    posting shares with the resume, then boosts for seniority alignment.
    Works for any career field.
    """
    if not resume_text.strip():
        return 50, ['no resume loaded — upload to enable scoring']

    score = 0
    reasons = []
    title_lower = job_title.lower()
    desc_lower  = (job_desc or '').lower()
    resume_lower = resume_text.lower()

    # ── 1. Tokenize for overlap scoring ─────────────────────────────────────
    resume_tokens = set(_tokenize(resume_lower))
    title_tokens  = set(_tokenize(title_lower))
    desc_tokens   = set(_tokenize(desc_lower))
    all_job_tokens = title_tokens | desc_tokens

    # ── 2. Title keyword overlap ─────────────────────────────────────────────
    title_overlap = title_tokens & resume_tokens
    if title_overlap:
        title_score = min(len(title_overlap) * 10, 30)
        score += title_score
        top = sorted(title_overlap, key=len, reverse=True)[:3]
        reasons.append(f"title match: {', '.join(top)}")

    # ── 3. Skills / abbreviation overlap ────────────────────────────────────
    skill_hits = [s for s in resume_skills if s in all_job_tokens and len(s) > 2]
    if skill_hits:
        score += min(len(skill_hits) * 4, 24)
        reasons.append(f"skills: {', '.join(skill_hits[:5])}")

    # ── 4. Seniority match ──────────────────────────────────────────────────
    senior_kw = ['senior', 'sr', 'lead', 'manager', 'director', 'head',
                 'principal', 'staff', 'vp', 'vice president', 'avp',
                 'associate director', 'program manager', 'supervisor']
    resume_is_senior = any(k in resume_lower for k in senior_kw)
    job_is_senior    = any(k in title_lower  for k in senior_kw)
    if resume_is_senior and job_is_senior:
        score += 15
        reasons.append('seniority match')
    elif resume_is_senior and not job_is_senior:
        score -= 5  # overqualified signal

    # ── 5. Description body overlap ──────────────────────────────────────────
    desc_overlap = desc_tokens & resume_tokens
    # Strip generic words that appear everywhere
    quality_overlap = [
        w for w in desc_overlap
        if len(w) > 4 and w not in {'other', 'their', 'these', 'those',
                                     'there', 'about', 'after', 'before',
                                     'under', 'over', 'between', 'through'}
    ]
    if quality_overlap:
        score += min(len(quality_overlap) * 2, 20)
        if len(reasons) < 4:
            top = sorted(quality_overlap, key=len, reverse=True)[:3]
            reasons.append(f"description overlap: {', '.join(top)}")

    return min(round(score, 1), 100), reasons

# ──────────────────────────────────────────────
# Job scraping / import
# ──────────────────────────────────────────────

def fetch_greenhouse_jobs(company_name, base_url):
    """Try to fetch jobs from Greenhouse API for a company."""
    # Derive slug from URL or fall back to lowercased name
    slug = company_name.lower().replace(' ', '').replace('-', '')
    # Try to extract from URL e.g. boards.greenhouse.io/SLUG
    import re as _re
    m = _re.search(r'greenhouse\.io/([^/?#]+)', base_url or '')
    if m:
        slug = m.group(1)

    try:
        url = f"https://boards-api.greenhouse.io/v1/boards/{slug}/jobs?content=true"
        resp = requests.get(url, timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
        if resp.status_code != 200:
            return []
        data = resp.json()
        jobs = []
        for j in data.get('jobs', [])[:150]:
            jobs.append({
                'title': j.get('title', ''),
                'company': company_name,
                'location': j.get('location', {}).get('name', ''),
                'url': j.get('absolute_url', ''),
                'source': 'greenhouse',
                'raw_description': BeautifulSoup(j.get('content', ''), 'html.parser').get_text()[:2000],
            })
        return jobs
    except Exception:
        return []


def fetch_lever_jobs(company_name, base_url):
    """Fetch jobs from Lever API for a company."""
    import re as _re
    slug = company_name.lower().replace(' ', '').replace('-', '')
    m = _re.search(r'jobs\.lever\.co/([^/?#]+)', base_url or '')
    if m:
        slug = m.group(1)

    try:
        url = f"https://api.lever.co/v0/postings/{slug}?mode=json"
        resp = requests.get(url, timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
        if resp.status_code != 200:
            return []
        data = resp.json()
        jobs = []
        for j in data[:150]:
            cats = j.get('categories', {})
            location = cats.get('location', '') or cats.get('allLocations', [''])[0] if j.get('categories') else ''
            desc_html = j.get('descriptionPlain', '') or j.get('description', '')
            jobs.append({
                'title': j.get('text', ''),
                'company': company_name,
                'location': location,
                'url': j.get('hostedUrl', ''),
                'source': 'lever',
                'raw_description': BeautifulSoup(desc_html, 'html.parser').get_text()[:2000] if '<' in desc_html else desc_html[:2000],
            })
        return jobs
    except Exception:
        return []


def fetch_ashby_jobs(company_name, base_url):
    """Fetch jobs from Ashby API for a company."""
    import re as _re
    slug = company_name.lower().replace(' ', '').replace('-', '')
    m = _re.search(r'jobs\.ashbyhq\.com/([^/?#]+)', base_url or '')
    if m:
        slug = m.group(1)

    try:
        url = f"https://api.ashbyhq.com/posting-api/job-board/{slug}"
        resp = requests.get(url, timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
        if resp.status_code != 200:
            return []
        data = resp.json()
        jobs = []
        for j in data.get('jobs', [])[:150]:
            loc = j.get('location', '') or (j.get('locationName') or '')
            jobs.append({
                'title': j.get('title', ''),
                'company': company_name,
                'location': loc,
                'url': j.get('jobUrl', ''),
                'source': 'ashby',
                'raw_description': BeautifulSoup(j.get('descriptionHtml', ''), 'html.parser').get_text()[:2000],
            })
        return jobs
    except Exception:
        return []

def fetch_remoteok_jobs(keyword_filter=None):
    """Fetch remote jobs from RemoteOK's public API (no auth required)."""
    try:
        url = "https://remoteok.com/api"
        resp = requests.get(url, timeout=15, headers={
            'User-Agent': 'Mozilla/5.0',
            'Accept': 'application/json'
        })
        if resp.status_code != 200:
            return []
        data = resp.json()
        jobs = []
        # First item is a legal notice object, skip it
        for j in data[1:]:
            if not isinstance(j, dict):
                continue
            title = j.get('position', '')
            company = j.get('company', '')
            if not title or not company:
                continue
            # Optional keyword filter — pass keyword_filter to narrow results
            tags = ' '.join(j.get('tags', [])).lower()
            combined = (title + ' ' + tags).lower()
            if keyword_filter and not any(t in combined for t in keyword_filter):
                continue
            jobs.append({
                'title': title,
                'company': company,
                'location': 'Remote',
                'url': j.get('url', ''),
                'source': 'remoteok',
                'remote': True,
                'raw_description': j.get('description', '')[:2000],
            })
        return jobs
    except Exception:
        return []


def scrape_page_for_jobs(url, company_name):
    """Basic scraper for direct career pages."""
    try:
        resp = requests.get(url, timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
        soup = BeautifulSoup(resp.text, 'html.parser')
        # Look for job-like links
        job_links = []
        for a in soup.find_all('a', href=True):
            text = a.get_text(strip=True)
            href = a['href']
            if len(text) > 10 and any(k in href.lower() for k in ['job', 'career', 'position', 'opening']):
                job_links.append({'title': text, 'url': href, 'company': company_name, 'source': 'scraped'})
        return job_links[:50]
    except Exception:
        return []

# ──────────────────────────────────────────────
# Routes - Static
# ──────────────────────────────────────────────

@app.route('/')
def index():
    return send_from_directory('templates', 'index.html')

@app.route('/static/<path:path>')
def static_files(path):
    return send_from_directory('static', path)

# ──────────────────────────────────────────────
# Routes - Resume
# ──────────────────────────────────────────────

@app.route('/api/resume', methods=['GET'])
def get_resume():
    text, skills = load_resume_text()
    return jsonify({'text': text[:500] + '...' if len(text) > 500 else text, 'skills': skills, 'loaded': bool(text)})

@app.route('/api/resume/upload', methods=['POST'])
def upload_resume():
    if 'file' not in request.files:
        # Try text body
        data = request.get_json(silent=True)
        if data and 'text' in data:
            text = data['text']
        else:
            return jsonify({'error': 'No file or text provided'}), 400
    else:
        f = request.files['file']
        filename = f.filename.lower()
        if filename.endswith('.docx'):
            from docx import Document
            import io
            doc = Document(io.BytesIO(f.read()))
            text = '\n'.join(p.text for p in doc.paragraphs)
        elif filename.endswith('.txt'):
            text = f.read().decode('utf-8')
        elif filename.endswith('.pdf'):
            import pdfplumber, io as _io
            pdf_bytes = f.read()
            with pdfplumber.open(_io.BytesIO(pdf_bytes)) as pdf:
                text = '\n'.join(
                    page.extract_text() or '' for page in pdf.pages
                )
            if not text.strip():
                return jsonify({'error': 'Could not extract text from PDF — try a .docx or .txt instead'}), 400
        else:
            text = f.read().decode('utf-8', errors='replace')

    skills = extract_skills_from_resume(text)
    conn = get_db()
    conn.execute("""
        INSERT OR REPLACE INTO resume_data (id, content, skills, last_updated)
        VALUES (1, ?, ?, ?)
    """, (text, json.dumps(skills), datetime.now().isoformat()))
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'skills': skills, 'char_count': len(text)})

# ──────────────────────────────────────────────
# Routes - Jobs CRUD
# ──────────────────────────────────────────────

@app.route('/api/jobs', methods=['GET'])
def get_jobs():
    status = request.args.get('status')
    search = request.args.get('q', '').lower()
    remote_only = request.args.get('remote') == '1'

    conn = get_db()
    query = "SELECT * FROM jobs WHERE 1=1"
    params = []
    if status:
        query += " AND status=?"
        params.append(status)
    if remote_only:
        query += " AND remote=1"
    if search:
        query += " AND (LOWER(title) LIKE ? OR LOWER(company) LIKE ? OR LOWER(notes) LIKE ?)"
        params += [f'%{search}%', f'%{search}%', f'%{search}%']
    query += " ORDER BY score DESC, date_added DESC"

    rows = conn.execute(query, params).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/jobs', methods=['POST'])
def add_job():
    data = request.get_json()
    resume_text, resume_skills = load_resume_text()
    score, reasons = score_job_against_resume(
        data.get('title', ''),
        data.get('raw_description', '') + ' ' + data.get('notes', ''),
        resume_text,
        resume_skills
    )

    conn = get_db()
    cur = conn.execute("""
        INSERT INTO jobs (title, company, location, remote, salary, url, source, status, score,
                          match_reasons, notes, date_added, date_updated, raw_description)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (
        data.get('title', ''),
        data.get('company', ''),
        data.get('location', ''),
        1 if data.get('remote') else 0,
        data.get('salary', ''),
        data.get('url', ''),
        data.get('source', 'manual'),
        data.get('status', 'interested'),
        data.get('score', score),
        json.dumps(data.get('match_reasons', reasons)),
        data.get('notes', ''),
        datetime.now().isoformat(),
        datetime.now().isoformat(),
        data.get('raw_description', ''),
    ))
    job_id = cur.lastrowid
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'id': job_id, 'score': score, 'reasons': reasons})

@app.route('/api/jobs/<int:job_id>', methods=['PUT'])
def update_job(job_id):
    data = request.get_json()
    conn = get_db()
    fields = []
    values = []
    allowed = ['title', 'company', 'location', 'remote', 'salary', 'url',
               'status', 'score', 'notes', 'date_applied', 'raw_description']
    for k in allowed:
        if k in data:
            fields.append(f"{k}=?")
            values.append(data[k])

    if not fields:
        return jsonify({'error': 'Nothing to update'}), 400

    fields.append("date_updated=?")
    values.append(datetime.now().isoformat())
    values.append(job_id)

    conn.execute(f"UPDATE jobs SET {', '.join(fields)} WHERE id=?", values)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/jobs/<int:job_id>', methods=['DELETE'])
def delete_job(job_id):
    conn = get_db()
    conn.execute("DELETE FROM jobs WHERE id=?", (job_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/jobs/import-csv', methods=['POST'])
def import_csv():
    """Import jobs from a CSV (matching the format of the uploaded file)."""
    import io, csv
    f = request.files.get('file')
    if not f:
        return jsonify({'error': 'No file'}), 400

    content = f.read().decode('utf-8')
    reader = csv.DictReader(io.StringIO(content))
    resume_text, resume_skills = load_resume_text()

    conn = get_db()
    imported = 0
    for row in reader:
        title = row.get('title', '').strip()
        company = row.get('company', '').strip()
        if not title or not company:
            continue

        # Check for dup
        existing = conn.execute(
            "SELECT id FROM jobs WHERE title=? AND company=?", (title, company)
        ).fetchone()
        if existing:
            continue

        # Use CSV score if available, else re-score
        csv_score = row.get('score', '')
        try:
            score = float(csv_score)
            reasons = [r.strip() for r in row.get('reasons', '').split('|') if r.strip()]
        except (ValueError, TypeError):
            score, reasons = score_job_against_resume(title, '', resume_text, resume_skills)

        location = row.get('location', '')
        remote_val = str(row.get('remote', '')).lower() in ('true', '1', 'yes', 'remote')

        conn.execute("""
            INSERT INTO jobs (title, company, location, remote, url, source, status,
                              score, match_reasons, date_added, date_updated)
            VALUES (?,?,?,?,?,?,?,?,?,?,?)
        """, (
            title, company, location, 1 if remote_val else 0,
            row.get('url', ''), row.get('source', 'csv'), 'interested',
            score, json.dumps(reasons),
            datetime.now().isoformat(), datetime.now().isoformat()
        ))
        imported += 1

    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'imported': imported})

# ──────────────────────────────────────────────
# Routes - Sources & Fetch
# ──────────────────────────────────────────────

@app.route('/api/jobs/export', methods=['GET'])
def export_jobs_csv():
    """Export all jobs as a CSV file with full pipeline status and actions."""
    import csv, io
    conn = get_db()
    rows = conn.execute("""
        SELECT title, company, location, remote, salary, url, source, status,
               score, match_reasons, notes, date_added, date_applied, date_updated
        FROM jobs
        ORDER BY score DESC, date_added DESC
    """).fetchall()
    conn.close()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        'Title', 'Company', 'Location', 'Remote', 'Salary', 'URL', 'Source',
        'Status', 'Match Score', 'Match Reasons', 'Competition Level', 'Notes',
        'Date Added', 'Date Applied', 'Last Updated'
    ])
    for r in rows:
        try:
            reasons = ', '.join(json.loads(r['match_reasons'] or '[]'))
        except Exception:
            reasons = r['match_reasons'] or ''
        competition = estimate_competition(r['title'], r['company'])
        writer.writerow([
            r['title'],
            r['company'],
            r['location'] or '',
            'Yes' if r['remote'] else 'No',
            r['salary'] or '',
            r['url'] or '',
            r['source'] or '',
            r['status'] or '',
            r['score'] or '',
            reasons,
            competition,
            r['notes'] or '',
            (r['date_added'] or '')[:10],
            (r['date_applied'] or '')[:10],
            (r['date_updated'] or '')[:10],
        ])

    from flask import Response
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': 'attachment; filename=jobtracker_export.csv'}
    )


@app.route('/api/sources', methods=['GET'])
def get_sources():
    conn = get_db()
    rows = conn.execute("SELECT * FROM sources ORDER BY active DESC, name").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/sources', methods=['POST'])
def add_source():
    data = request.get_json()
    conn = get_db()
    cur = conn.execute(
        "INSERT INTO sources (name, url, type, active) VALUES (?,?,?,1)",
        (data.get('name',''), data.get('url',''), data.get('type','direct'))
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'id': cur.lastrowid})

@app.route('/api/sources/<int:src_id>', methods=['PUT'])
def update_source(src_id):
    data = request.get_json()
    conn = get_db()
    conn.execute(
        "UPDATE sources SET name=?, url=?, type=?, active=? WHERE id=?",
        (data['name'], data['url'], data.get('type','direct'), data.get('active',1), src_id)
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/sources/<int:src_id>', methods=['DELETE'])
def delete_source(src_id):
    conn = get_db()
    conn.execute("DELETE FROM sources WHERE id=?", (src_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/sources/<int:src_id>/fetch', methods=['POST'])
def fetch_source(src_id):
    conn = get_db()
    source = conn.execute("SELECT * FROM sources WHERE id=?", (src_id,)).fetchone()
    if not source:
        conn.close()
        return jsonify({'error': 'Source not found'}), 404

    source = dict(source)
    resume_text, resume_skills = load_resume_text()
    jobs = []

    if source['type'] == 'greenhouse':
        jobs = fetch_greenhouse_jobs(source['name'], source['url'])
    elif source['type'] == 'lever':
        jobs = fetch_lever_jobs(source['name'], source['url'])
    elif source['type'] == 'ashby':
        jobs = fetch_ashby_jobs(source['name'], source['url'])
    elif source['type'] == 'remoteok':
        jobs = fetch_remoteok_jobs()
    else:
        jobs = scrape_page_for_jobs(source['url'], source['name'])

    added = 0
    for j in jobs:
        existing = conn.execute(
            "SELECT id FROM jobs WHERE title=? AND company=?",
            (j['title'], j['company'])
        ).fetchone()
        if existing:
            continue

        score, reasons = score_job_against_resume(
            j['title'], j.get('raw_description', ''), resume_text, resume_skills
        )

        location = j.get('location', '')
        remote_val = any(k in location.lower() for k in ['remote', 'anywhere'])

        conn.execute("""
            INSERT INTO jobs (title, company, location, remote, url, source, status,
                              score, match_reasons, date_added, date_updated, raw_description)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
        """, (
            j['title'], j['company'], location, 1 if remote_val else 0,
            j.get('url',''), j.get('source', source['type']), 'interested',
            score, json.dumps(reasons),
            datetime.now().isoformat(), datetime.now().isoformat(),
            j.get('raw_description','')[:3000]
        ))
        added += 1

    conn.execute(
        "UPDATE sources SET last_fetched=?, job_count=? WHERE id=?",
        (datetime.now().isoformat(), len(jobs), src_id)
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'fetched': len(jobs), 'added': added})

# ──────────────────────────────────────────────
# Routes - Stats
# ──────────────────────────────────────────────

@app.route('/api/jobs/rescore', methods=['POST'])
def rescore_all_jobs():
    """Re-run the scoring algorithm on every job in the DB using the current resume."""
    resume_text, resume_skills = load_resume_text()
    if not resume_text:
        return jsonify({'error': 'No resume loaded — upload your resume first'}), 400

    conn = get_db()
    rows = conn.execute("SELECT id, title, raw_description, notes FROM jobs").fetchall()
    updated = 0
    zeroed  = 0
    for r in rows:
        score, reasons = score_job_against_resume(
            r['title'],
            (r['raw_description'] or '') + ' ' + (r['notes'] or ''),
            resume_text,
            resume_skills
        )
        conn.execute(
            "UPDATE jobs SET score=?, match_reasons=?, date_updated=? WHERE id=?",
            (score, json.dumps(reasons), datetime.now().isoformat(), r['id'])
        )
        updated += 1
        if score == 0:
            zeroed += 1

    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'rescored': updated, 'zeroed_out': zeroed})


@app.route('/api/stats', methods=['GET'])
def get_stats():
    conn = get_db()
    total = conn.execute("SELECT COUNT(*) FROM jobs").fetchone()[0]
    by_status = conn.execute(
        "SELECT status, COUNT(*) as n FROM jobs GROUP BY status"
    ).fetchall()
    top_companies = conn.execute(
        "SELECT company, COUNT(*) as n FROM jobs GROUP BY company ORDER BY n DESC LIMIT 5"
    ).fetchall()
    conn.close()
    return jsonify({
        'total': total,
        'by_status': {r['status']: r['n'] for r in by_status},
        'top_companies': [{'company': r['company'], 'count': r['n']} for r in top_companies],
    })

# ──────────────────────────────────────────────
# URL scrape helper
# ──────────────────────────────────────────────

@app.route('/api/scrape-url', methods=['POST'])
def scrape_url():
    """Scrape a job posting URL and return structured data."""
    data = request.get_json()
    url = data.get('url', '')
    if not url:
        return jsonify({'error': 'No URL'}), 400

    try:
        resp = requests.get(url, timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
        soup = BeautifulSoup(resp.text, 'html.parser')

        # Remove scripts/styles
        for tag in soup(['script', 'style', 'nav', 'footer']):
            tag.decompose()

        text = soup.get_text(separator='\n', strip=True)[:5000]

        # Try to extract title
        title = ''
        for tag in ['h1', 'h2']:
            el = soup.find(tag)
            if el:
                title = el.get_text(strip=True)
                break

        return jsonify({'ok': True, 'title': title, 'description': text})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/detect-ats', methods=['POST'])
def detect_ats():
    """
    Given a company name or URL, probe Greenhouse/Lever/Ashby APIs
    to detect which ATS the company uses.
    """
    data = request.get_json()
    name = data.get('name', '').strip()
    url  = data.get('url', '').strip()

    # Extract slug from URL if provided
    import re as _re
    slug = name.lower().replace(' ', '').replace('-', '').replace('.', '')

    # Override slug from known URL patterns
    for pattern, ats in [
        (r'greenhouse\.io/([^/?#]+)', 'greenhouse'),
        (r'lever\.co/([^/?#]+)', 'lever'),
        (r'ashbyhq\.com/([^/?#]+)', 'ashby'),
        (r'remoteok\.com', 'remoteok'),
    ]:
        m = _re.search(pattern, url)
        if m:
            return jsonify({'ats': ats, 'slug': m.group(1), 'url': url})

    # Probe each ATS
    probes = [
        ('greenhouse', f"https://boards-api.greenhouse.io/v1/boards/{slug}/jobs"),
        ('lever',      f"https://api.lever.co/v0/postings/{slug}?mode=json"),
        ('ashby',      f"https://api.ashbyhq.com/posting-api/job-board/{slug}"),
    ]
    for ats, probe_url in probes:
        try:
            r = requests.head(probe_url, timeout=6, headers={'User-Agent': 'Mozilla/5.0'})
            if r.status_code == 200:
                return jsonify({'ats': ats, 'slug': slug, 'url': probe_url})
        except Exception:
            continue

    return jsonify({'ats': 'unknown', 'slug': slug})


@app.route('/api/smart-fetch', methods=['POST'])
def smart_fetch():
    """
    Given any URL, domain, or company name:
    1. Derive candidate slugs
    2. Probe Greenhouse / Lever / Ashby to detect ATS
    3. Fetch & score matching jobs
    4. Upsert the source into the DB
    Returns detected ATS, source id, and scored jobs.
    """
    data   = request.get_json()
    raw    = (data.get('input', '') or '').strip()
    if not raw:
        return jsonify({'error': 'No input provided'}), 400

    # ── Derive company name and slug candidates ──────────────────────────────
    import re as _re
    from urllib.parse import urlparse

    parsed = urlparse(raw if '://' in raw else 'https://' + raw)
    hostname = parsed.hostname or ''          # e.g. "boards.greenhouse.io" or "stripe.com"
    path_parts = [p for p in parsed.path.strip('/').split('/') if p]

    # Best-guess company name: last meaningful path segment or subdomain
    # e.g. jobs.lever.co/mercury → "mercury",  stripe.com → "stripe"
    def clean(s):
        return _re.sub(r'[^a-z0-9]', '', s.lower())

    slug_candidates = []

    # Known ATS host patterns — extract slug from path
    known_ats = {
        'boards.greenhouse.io': 'greenhouse',
        'greenhouse.io':        'greenhouse',
        'jobs.lever.co':        'lever',
        'lever.co':             'lever',
        'jobs.ashbyhq.com':     'ashby',
        'ashbyhq.com':          'ashby',
    }
    detected_ats = None
    for host_pat, ats_name in known_ats.items():
        if host_pat in hostname:
            detected_ats = ats_name
            if path_parts:
                slug_candidates.insert(0, clean(path_parts[0]))
            break

    # Derive from domain (strip common TLDs/prefixes)
    domain_base = _re.sub(r'^(www|jobs|careers|app)\.', '', hostname)
    domain_base = _re.sub(r'\.(com|io|co|net|org|ai|app)(\..*)?$', '', domain_base)
    slug_candidates.append(clean(domain_base))

    # Also try the raw input itself as a slug
    slug_candidates.append(clean(_re.sub(r'https?://', '', raw).split('/')[0].split('.')[0]))

    # Deduplicate, drop empties
    seen = set()
    slugs = []
    for s in slug_candidates:
        if s and s not in seen:
            seen.add(s)
            slugs.append(s)

    # Company display name: title-case the best slug
    company_name = data.get('name') or domain_base.replace('-', ' ').title() or slugs[0].title()

    # ── Probe each ATS with each slug ────────────────────────────────────────
    ats_probes = [
        ('greenhouse', lambda sl: f"https://boards-api.greenhouse.io/v1/boards/{sl}/jobs"),
        ('lever',      lambda sl: f"https://api.lever.co/v0/postings/{sl}?mode=json"),
        ('ashby',      lambda sl: f"https://api.ashbyhq.com/posting-api/job-board/{sl}"),
    ]

    found_ats   = detected_ats
    found_slug  = slugs[0] if slugs else ''
    found_url   = raw

    if not found_ats:
        for ats_name, url_fn in ats_probes:
            for sl in slugs:
                try:
                    r = requests.head(url_fn(sl), timeout=6,
                                      headers={'User-Agent': 'Mozilla/5.0'})
                    if r.status_code == 200:
                        found_ats  = ats_name
                        found_slug = sl
                        found_url  = url_fn(sl)
                        break
                except Exception:
                    continue
            if found_ats:
                break

    # ── Fetch jobs ───────────────────────────────────────────────────────────
    jobs_raw = []
    if found_ats == 'greenhouse':
        jobs_raw = fetch_greenhouse_jobs(company_name, found_url)
        if not jobs_raw and found_slug != slugs[0]:
            jobs_raw = fetch_greenhouse_jobs(company_name,
                           f"https://boards.greenhouse.io/{found_slug}")
    elif found_ats == 'lever':
        jobs_raw = fetch_lever_jobs(company_name, found_url)
    elif found_ats == 'ashby':
        jobs_raw = fetch_ashby_jobs(company_name, found_url)
    else:
        # Fall back: try to discover careers page and scrape
        career_paths = ['/careers', '/jobs', '/join', '/work-with-us',
                        '/about/careers', '/company/careers', '/en/careers']
        base = f"https://{hostname}" if hostname else raw
        for cp in career_paths:
            try:
                r = requests.get(base + cp, timeout=8,
                                 headers={'User-Agent': 'Mozilla/5.0'}, allow_redirects=True)
                if r.status_code == 200:
                    found_url = r.url
                    jobs_raw  = scrape_page_for_jobs(found_url, company_name)
                    found_ats = 'direct'
                    break
            except Exception:
                continue

    # ── Score & filter ───────────────────────────────────────────────────────
    resume_text, resume_skills = load_resume_text()
    conn = get_db()

    # Upsert source
    src_row = conn.execute("SELECT id FROM sources WHERE name=?", (company_name,)).fetchone()
    if src_row:
        src_id = src_row['id']
    else:
        cur = conn.execute(
            "INSERT INTO sources (name, url, type, active) VALUES (?,?,?,1)",
            (company_name, found_url, found_ats or 'direct')
        )
        src_id = cur.lastrowid

    added = 0
    results = []
    for j in jobs_raw:
        score, reasons = score_job_against_resume(
            j['title'], j.get('raw_description', ''), resume_text, resume_skills
        )
        location  = j.get('location', '')
        remote_val = any(k in location.lower() for k in ['remote', 'anywhere']) or j.get('remote', False)

        existing = conn.execute(
            "SELECT id FROM jobs WHERE title=? AND company=?",
            (j['title'], j['company'])
        ).fetchone()
        if not existing:
            conn.execute("""
                INSERT INTO jobs (title, company, location, remote, url, source, status,
                                  score, match_reasons, date_added, date_updated, raw_description)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
            """, (
                j['title'], j['company'], location, 1 if remote_val else 0,
                j.get('url', ''), found_ats or 'direct', 'interested',
                score, json.dumps(reasons),
                datetime.now().isoformat(), datetime.now().isoformat(),
                j.get('raw_description', '')[:3000]
            ))
            added += 1

        results.append({
            'title':   j['title'],
            'company': j['company'],
            'score':   score,
            'reasons': reasons,
            'url':     j.get('url', ''),
        })

    conn.execute("UPDATE sources SET last_fetched=?, job_count=? WHERE id=?",
                 (datetime.now().isoformat(), len(jobs_raw), src_id))
    conn.commit()
    conn.close()

    # Sort by score desc for the preview
    results.sort(key=lambda x: x['score'], reverse=True)

    return jsonify({
        'ok':          True,
        'ats':         found_ats or 'unknown',
        'company':     company_name,
        'source_id':   src_id,
        'fetched':     len(jobs_raw),
        'added':       added,
        'top_matches': results[:10],
    })


@app.route('/api/jobboard-search-urls', methods=['GET'])
def jobboard_search_urls():
    """
    Return pre-built search URLs for major job boards.
    Pass ?q=your+search+terms to customise the search query,
    or the endpoint will use terms derived from your loaded resume skills.
    """
    # Honour an explicit query param; otherwise derive from resume
    q = request.args.get('q', '').strip()
    if not q:
        _, resume_skills = load_resume_text()
        # Use up to 4 abbreviation-style skills (ALL-CAPS) as search terms — most precise
        abbrev_skills = [s.upper() for s in resume_skills if s.isupper() or (s.upper() == s and len(s) <= 5)]
        word_skills   = [s for s in resume_skills if not (s.upper() == s and len(s) <= 5)]
        top_terms = (abbrev_skills[:3] + word_skills[:2]) or ['jobs']
        q = ' '.join(top_terms[:4])

    eq = requests.utils.quote(q)
    eq_plus = q.replace(' ', '+')

    boards = [
        {
            'name': 'LinkedIn Jobs',
            'icon': 'in',
            'color': '#0a66c2',
            'url': f'https://www.linkedin.com/jobs/search/?keywords={eq_plus}',
            'note': 'Add Remote filter in LinkedIn UI'
        },
        {
            'name': 'Indeed',
            'icon': 'in',
            'color': '#2164f3',
            'url': f'https://www.indeed.com/jobs?q={eq_plus}&l=Remote',
            'note': 'Remote filtered'
        },
        {
            'name': 'Glassdoor',
            'icon': 'gd',
            'color': '#0caa41',
            'url': f'https://www.glassdoor.com/Job/jobs.htm?sc.keyword={eq}&locT=N&remoteWorkType=1',
            'note': 'Remote filtered'
        },
        {
            'name': 'Wellfound (AngelList)',
            'icon': 'wf',
            'color': '#000',
            'url': f'https://wellfound.com/jobs?q={eq_plus}&remote=true',
            'note': 'Startup-focused'
        },
        {
            'name': 'Built In',
            'icon': 'bi',
            'color': '#5c36f3',
            'url': f'https://builtin.com/jobs?search={eq_plus}&remote=true',
            'note': 'Tech company focused'
        },
        {
            'name': 'ZipRecruiter',
            'icon': 'zr',
            'color': '#4a90d9',
            'url': f'https://www.ziprecruiter.com/jobs-search?search={eq_plus}&location=Remote',
            'note': 'Remote filtered'
        },
        {
            'name': 'FlexJobs',
            'icon': 'fj',
            'color': '#00b373',
            'url': f'https://www.flexjobs.com/search?search={eq_plus}&location=Remote',
            'note': 'Vetted remote-only'
        },
        {
            'name': 'RemoteOK',
            'icon': 'ro',
            'color': '#00d4aa',
            'url': f'https://remoteok.com/remote-{eq_plus.replace("+", "-")}-jobs',
            'note': 'Remote-only board'
        },
        {
            'name': 'USAJobs',
            'icon': 'us',
            'color': '#1a3c6e',
            'url': f'https://www.usajobs.gov/search/results/?k={eq_plus}&p=1',
            'note': 'Federal / government roles'
        },
    ]
    return jsonify(boards)



if __name__ == '__main__':
    init_db()

    # Auto-import a resume from the data/ folder if not already loaded
    resume_text, _ = load_resume_text()
    if not resume_text:
        data_dir = os.path.join(os.path.dirname(__file__), 'data')
        loaded = False
        # Check for resume files in order of preference
        candidates = [
            ('resume.pdf',  'pdf'),
            ('resume.docx', 'docx'),
            ('resume.txt',  'txt'),
        ]
        for fname, ftype in candidates:
            fpath = os.path.join(data_dir, fname)
            if not os.path.exists(fpath):
                continue
            try:
                if ftype == 'docx':
                    from docx import Document
                    doc = Document(fpath)
                    text = '\n'.join(p.text for p in doc.paragraphs)
                elif ftype == 'pdf':
                    import pdfplumber
                    with pdfplumber.open(fpath) as pdf:
                        text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
                else:
                    with open(fpath, encoding='utf-8') as fh:
                        text = fh.read()

                if text.strip():
                    skills = extract_skills_from_resume(text)
                    conn = get_db()
                    conn.execute(
                        "INSERT OR REPLACE INTO resume_data (id, content, skills, last_updated) VALUES (1,?,?,?)",
                        (text, json.dumps(skills), datetime.now().isoformat())
                    )
                    conn.commit()
                    conn.close()
                    print(f"✓ Resume auto-loaded from {fname} ({len(text)} chars, {len(skills)} skills)")
                    loaded = True
                    break
            except Exception as e:
                print(f"Resume auto-load ({fname}) skipped: {e}")
        if not loaded:
            print("ℹ  No resume found in data/ — upload via the UI to enable job scoring")

    print("\n🚀 JobTracker running at http://localhost:5000\n")
    app.run(debug=True, host='0.0.0.0', port=5001)
