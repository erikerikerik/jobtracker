"""
job-agent-web — Flask web interface for job-agent.
Run via:  job-agent-web
Then open: http://localhost:5001
"""

from __future__ import annotations

import csv
import io
import json
import os
import re
import sqlite3
from collections import Counter
from datetime import datetime

from bs4 import BeautifulSoup
from flask import Flask, Response, jsonify, request, send_from_directory
import requests

# ── Paths ──────────────────────────────────────────────────────────────────────
_HERE = os.path.dirname(os.path.abspath(__file__))
_TEMPLATE_DIR = os.path.join(_HERE, "templates")

app = Flask(__name__, template_folder=_TEMPLATE_DIR)

# Data lives next to wherever the user runs the server (i.e. CWD/data)
DB_PATH      = os.path.join("data", "jobs.db")
PROFILE_PATH = os.path.join("data", "search_profile.json")

DEFAULT_SEARCH_PROFILE = {
    "target_titles": [],
    "preferred_keywords": [],
    "preferred_industries": [],
    "preferred_seniority": [],
    "hard_exclusions": [],
    "minimum_keyword_matches": 0,
    "minimum_industry_matches": 0,
    "require_title_or_keyword": False,
    "require_industry": False,
    "require_seniority": False,
}

# ──────────────────────────────────────────────────────────────────────────────
# Database
# ──────────────────────────────────────────────────────────────────────────────

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conn = get_db()
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS jobs (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            title           TEXT NOT NULL,
            company         TEXT NOT NULL,
            location        TEXT,
            remote          INTEGER DEFAULT 0,
            salary          TEXT,
            url             TEXT,
            source          TEXT DEFAULT 'manual',
            status          TEXT DEFAULT 'interested',
            score           REAL,
            match_reasons   TEXT,
            notes           TEXT,
            date_added      TEXT,
            date_applied    TEXT,
            date_updated    TEXT,
            raw_description TEXT,
            is_new          INTEGER DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS sources (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            name         TEXT NOT NULL,
            url          TEXT NOT NULL,
            type         TEXT DEFAULT 'greenhouse',
            active       INTEGER DEFAULT 1,
            last_fetched TEXT,
            job_count    INTEGER DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS resume_data (
            id           INTEGER PRIMARY KEY,
            content      TEXT,
            skills       TEXT,
            last_updated TEXT
        );
    """)
    conn.commit()

    # Add is_new column to existing DBs that pre-date this migration
    try:
        conn.execute("ALTER TABLE jobs ADD COLUMN is_new INTEGER DEFAULT 0")
        conn.commit()
    except Exception:
        pass  # column already exists

    # Seed default sources if none present
    default_sources = [
        ("Stripe",    "https://boards.greenhouse.io/stripe",    "greenhouse", 1),
        ("Brex",      "https://boards.greenhouse.io/brex",      "greenhouse", 1),
        ("Ramp",      "https://boards.greenhouse.io/ramp",      "greenhouse", 1),
        ("Plaid",     "https://boards.greenhouse.io/plaid",     "greenhouse", 1),
        ("Airbnb",    "https://boards.greenhouse.io/airbnb",    "greenhouse", 1),
        ("Mercury",   "https://jobs.lever.co/mercury",          "lever",      1),
        ("RemoteOK",  "https://remoteok.com/api",               "remoteok",   1),
        ("Remotive",  "https://remotive.com/api/remote-jobs",   "remotive",   1),
    ]
    for name, url, type_, active in default_sources:
        if not conn.execute("SELECT id FROM sources WHERE name=?", (name,)).fetchone():
            conn.execute(
                "INSERT INTO sources (name, url, type, active) VALUES (?,?,?,?)",
                (name, url, type_, active),
            )
    conn.commit()
    conn.close()


# ──────────────────────────────────────────────────────────────────────────────
# Resume helpers
# ──────────────────────────────────────────────────────────────────────────────

def load_resume_text():
    conn = get_db()
    row = conn.execute("SELECT content, skills FROM resume_data WHERE id=1").fetchone()
    conn.close()
    if row:
        return row["content"], json.loads(row["skills"] or "[]")
    return "", []


_STOPWORDS = {
    "a", "an", "the", "and", "or", "but", "in", "on", "at", "to", "for",
    "of", "with", "by", "from", "up", "about", "into", "through", "during",
    "is", "are", "was", "were", "be", "been", "being", "have", "has", "had",
    "do", "does", "did", "will", "would", "could", "should", "may", "might",
    "must", "can", "this", "that", "these", "those", "i", "you", "he", "she",
    "we", "they", "it", "as", "if", "than", "so", "yet", "both", "either",
    "not", "no", "nor", "too", "very", "just", "our", "your", "their", "its",
    "my", "his", "her", "who", "which", "what", "when", "where", "how", "all",
    "any", "each", "every", "more", "most", "other", "some", "such", "only",
    "also", "well", "new", "good", "high", "including", "across", "within",
    "experience", "skills", "ability", "strong", "excellent", "years", "year",
    "team", "role", "responsibilities", "qualifications", "required", "preferred",
    "minimum", "position", "candidate", "company", "business", "using",
    "ensure", "support", "provide", "develop", "manage", "working", "related",
    "relevant", "job", "opportunity", "looking", "join", "based",
}


def _tokenize(text: str) -> list[str]:
    words = re.findall(r"\b[a-zA-Z][a-zA-Z0-9+#\-\.]*\b", text.lower())
    return [w for w in words if len(w) > 2 and w not in _STOPWORDS]


def extract_skills_from_resume(text: str) -> list[str]:
    abbrevs = re.findall(r"\b[A-Z]{2,6}\b", text)
    tokens = _tokenize(text)
    counts = Counter(tokens)
    recurring = {w for w, c in counts.items() if c >= 2 and len(w) > 3}
    skills = set(a.lower() for a in abbrevs if len(a) >= 2)
    skills.update(recurring)
    noise = {
        "work", "time", "also", "make", "used", "use", "like", "one",
        "two", "three", "four", "five", "including", "within", "across",
        "other", "they", "that", "this", "with", "from", "have", "been",
    }
    skills -= noise
    return sorted(list(skills))[:100]


# ──────────────────────────────────────────────────────────────────────────────
# Search profile
# ──────────────────────────────────────────────────────────────────────────────

def _clean_profile_list(items):
    seen, cleaned = set(), []
    for item in items or []:
        text = str(item).strip().lower()
        if text and text not in seen:
            cleaned.append(text)
            seen.add(text)
    return cleaned


def load_search_profile():
    profile = dict(DEFAULT_SEARCH_PROFILE)
    if os.path.exists(PROFILE_PATH):
        try:
            with open(PROFILE_PATH, encoding="utf-8") as fh:
                stored = json.load(fh)
            if isinstance(stored, dict):
                profile.update(stored)
        except Exception:
            pass
    for key in ("target_titles", "preferred_keywords", "preferred_industries",
                "preferred_seniority", "hard_exclusions"):
        profile[key] = _clean_profile_list(profile.get(key, []))
    profile["minimum_keyword_matches"] = max(0, int(profile.get("minimum_keyword_matches") or 0))
    profile["minimum_industry_matches"] = max(0, int(profile.get("minimum_industry_matches") or 0))
    for key in ("require_title_or_keyword", "require_industry", "require_seniority"):
        profile[key] = bool(profile.get(key))
    return profile


def save_search_profile(profile):
    os.makedirs(os.path.dirname(PROFILE_PATH), exist_ok=True)
    normalized = dict(DEFAULT_SEARCH_PROFILE)
    normalized.update(profile or {})
    for key in ("target_titles", "preferred_keywords", "preferred_industries",
                "preferred_seniority", "hard_exclusions"):
        normalized[key] = _clean_profile_list(normalized.get(key, []))
    normalized["minimum_keyword_matches"] = max(0, int(normalized.get("minimum_keyword_matches") or 0))
    normalized["minimum_industry_matches"] = max(0, int(normalized.get("minimum_industry_matches") or 0))
    for key in ("require_title_or_keyword", "require_industry", "require_seniority"):
        normalized[key] = bool(normalized.get(key))
    with open(PROFILE_PATH, "w", encoding="utf-8") as fh:
        json.dump(normalized, fh, indent=2)
    return normalized


# ──────────────────────────────────────────────────────────────────────────────
# Scoring
# ──────────────────────────────────────────────────────────────────────────────

def estimate_competition(job_title: str, company: str) -> str:
    title_lower   = job_title.lower()
    company_lower = company.lower()
    high_brand = [
        "google", "apple", "meta", "amazon", "microsoft", "netflix", "spotify",
        "stripe", "airbnb", "uber", "lyft", "coinbase", "robinhood", "openai",
        "slack", "shopify", "twitter", "instagram", "linkedin", "salesforce",
    ]
    consulting = [
        "mckinsey", "deloitte", "kpmg", "pwc", "accenture", "bain",
        "boston consulting", "ernst", "guidehouse", "treliant",
    ]
    if any(b in company_lower for b in high_brand):
        return "🔴 HIGH (~500–1,000+ applicants)"
    if any(b in company_lower for b in consulting):
        return "🟡 MEDIUM (~200–400 applicants)"
    if any(k in title_lower for k in ["vp ", "vice president", "director", "head of", "principal", "avp"]):
        return "🟡 MEDIUM (~200–350 applicants)"
    if len(company) < 12:
        return "🟢 MEDIUM-LOW (~100–250 applicants)"
    return "🟡 MEDIUM (~150–300 applicants)"


def _phrase_hits(text: str, phrases: list[str]) -> list[str]:
    text = text.lower()
    return [p for p in phrases or [] if p and p in text]


def score_job_against_resume(
    job_title: str,
    job_desc: str,
    resume_text: str,
    resume_skills: list[str],
    search_profile: dict | None = None,
) -> tuple[float, list[str]]:
    search_profile = search_profile or load_search_profile()
    if not resume_text.strip():
        return 50, ["no resume loaded — upload to enable scoring"]

    score: float = 0
    reasons: list[str] = []
    title_lower   = job_title.lower()
    desc_lower    = (job_desc or "").lower()
    resume_lower  = resume_text.lower()
    full_job_text = f"{title_lower} {desc_lower}"

    hard_exclusions = _phrase_hits(full_job_text, search_profile.get("hard_exclusions", []))
    if hard_exclusions:
        return 0, [f"filtered: hard exclusion ({hard_exclusions[0]})"]

    resume_tokens  = set(_tokenize(resume_lower))
    title_tokens   = set(_tokenize(title_lower))
    desc_tokens    = set(_tokenize(desc_lower))
    all_job_tokens = title_tokens | desc_tokens

    target_title_hits = _phrase_hits(title_lower,   search_profile.get("target_titles", []))
    preferred_hits    = _phrase_hits(full_job_text, search_profile.get("preferred_keywords", []))
    industry_hits     = _phrase_hits(full_job_text, search_profile.get("preferred_industries", []))
    seniority_hits    = _phrase_hits(title_lower,   search_profile.get("preferred_seniority", []))

    if search_profile.get("require_title_or_keyword") and not target_title_hits:
        if len(preferred_hits) < search_profile.get("minimum_keyword_matches", 0):
            return 0, ["filtered: missing target title or required keywords"]
    if search_profile.get("require_industry"):
        if len(industry_hits) < search_profile.get("minimum_industry_matches", 0):
            return 0, ["filtered: missing preferred industry"]
    if search_profile.get("require_seniority") and not seniority_hits:
        return 0, ["filtered: missing preferred seniority"]

    if target_title_hits:
        score += min(len(target_title_hits) * 18, 36)
        reasons.append(f"profile title: {', '.join(target_title_hits[:3])}")
    if preferred_hits:
        score += min(len(preferred_hits) * 5, 30)
        reasons.append(f"profile keywords: {', '.join(preferred_hits[:4])}")
    if industry_hits:
        score += min(len(industry_hits) * 7, 21)
        reasons.append(f"industry: {', '.join(industry_hits[:3])}")
    if seniority_hits:
        score += min(len(seniority_hits) * 10, 20)
        reasons.append(f"seniority: {', '.join(seniority_hits[:2])}")

    title_overlap = title_tokens & resume_tokens
    if title_overlap:
        score += min(len(title_overlap) * 6, 18)
        top = sorted(title_overlap, key=len, reverse=True)[:3]
        if len(reasons) < 5:
            reasons.append(f"resume title overlap: {', '.join(top)}")

    skill_hits = [s for s in resume_skills if s in all_job_tokens and len(s) > 2]
    if skill_hits:
        score += min(len(skill_hits) * 3, 18)
        if len(reasons) < 5:
            reasons.append(f"resume skills: {', '.join(skill_hits[:5])}")

    senior_kw = [
        "senior", "sr", "lead", "manager", "director", "head",
        "principal", "staff", "vp", "vice president", "avp",
        "associate director", "program manager", "supervisor",
    ]
    resume_is_senior = any(k in resume_lower for k in senior_kw)
    job_is_senior    = any(k in title_lower  for k in senior_kw)
    if resume_is_senior and job_is_senior:
        score += 8
        if len(reasons) < 5:
            reasons.append("resume seniority match")
    elif resume_is_senior and not job_is_senior:
        score -= 5

    desc_overlap    = desc_tokens & resume_tokens
    quality_overlap = [
        w for w in desc_overlap
        if len(w) > 4 and w not in {"other", "their", "these", "those",
                                     "there", "about", "after", "before",
                                     "under", "over",  "between", "through"}
    ]
    if quality_overlap:
        score += min(len(quality_overlap) * 1.2, 12)
        if len(reasons) < 5:
            top = sorted(quality_overlap, key=len, reverse=True)[:3]
            reasons.append(f"description overlap: {', '.join(top)}")

    return min(round(score, 1), 100), reasons


def is_filtered_match(score: float, reasons: list[str]) -> bool:
    return score <= 0 and any(str(r).startswith("filtered:") for r in reasons or [])


# ──────────────────────────────────────────────────────────────────────────────
# Job source fetchers
# ──────────────────────────────────────────────────────────────────────────────

def _slug_from_url(pattern: str, url: str, fallback: str) -> str:
    m = re.search(pattern, url or "")
    return m.group(1) if m else fallback


def fetch_greenhouse_jobs(company_name: str, base_url: str) -> list[dict]:
    slug = _slug_from_url(r"greenhouse\.io/([^/?#]+)", base_url,
                          company_name.lower().replace(" ", "").replace("-", ""))
    try:
        resp = requests.get(
            f"https://boards-api.greenhouse.io/v1/boards/{slug}/jobs?content=true",
            timeout=10, headers={"User-Agent": "Mozilla/5.0"},
        )
        if resp.status_code != 200:
            return []
        return [
            {
                "title":           j.get("title", ""),
                "company":         company_name,
                "location":        j.get("location", {}).get("name", ""),
                "url":             j.get("absolute_url", ""),
                "source":          "greenhouse",
                "raw_description": BeautifulSoup(j.get("content", ""), "html.parser").get_text()[:2000],
            }
            for j in resp.json().get("jobs", [])[:150]
        ]
    except Exception:
        return []


def fetch_lever_jobs(company_name: str, base_url: str) -> list[dict]:
    slug = _slug_from_url(r"jobs\.lever\.co/([^/?#]+)", base_url,
                          company_name.lower().replace(" ", "").replace("-", ""))
    try:
        resp = requests.get(
            f"https://api.lever.co/v0/postings/{slug}?mode=json",
            timeout=10, headers={"User-Agent": "Mozilla/5.0"},
        )
        if resp.status_code != 200:
            return []
        jobs = []
        for j in resp.json()[:150]:
            cats     = j.get("categories", {})
            location = cats.get("location", "") if cats else ""
            desc_html = j.get("descriptionPlain", "") or j.get("description", "")
            jobs.append({
                "title":           j.get("text", ""),
                "company":         company_name,
                "location":        location,
                "url":             j.get("hostedUrl", ""),
                "source":          "lever",
                "raw_description": (
                    BeautifulSoup(desc_html, "html.parser").get_text()[:2000]
                    if "<" in desc_html else desc_html[:2000]
                ),
            })
        return jobs
    except Exception:
        return []


def fetch_ashby_jobs(company_name: str, base_url: str) -> list[dict]:
    slug = _slug_from_url(r"jobs\.ashbyhq\.com/([^/?#]+)", base_url,
                          company_name.lower().replace(" ", "").replace("-", ""))
    try:
        resp = requests.get(
            f"https://api.ashbyhq.com/posting-api/job-board/{slug}",
            timeout=10, headers={"User-Agent": "Mozilla/5.0"},
        )
        if resp.status_code != 200:
            return []
        return [
            {
                "title":           j.get("title", ""),
                "company":         company_name,
                "location":        j.get("location", "") or j.get("locationName", ""),
                "url":             j.get("jobUrl", ""),
                "source":          "ashby",
                "raw_description": BeautifulSoup(
                    j.get("descriptionHtml", ""), "html.parser"
                ).get_text()[:2000],
            }
            for j in resp.json().get("jobs", [])[:150]
        ]
    except Exception:
        return []


def fetch_remoteok_jobs(keyword_filter: list[str] | None = None) -> list[dict]:
    try:
        resp = requests.get(
            "https://remoteok.com/api",
            timeout=15,
            headers={"User-Agent": "Mozilla/5.0", "Accept": "application/json"},
        )
        if resp.status_code != 200:
            return []
        jobs = []
        for j in resp.json()[1:]:
            if not isinstance(j, dict):
                continue
            title   = j.get("position", "")
            company = j.get("company", "")
            if not title or not company:
                continue
            tags     = " ".join(j.get("tags", [])).lower()
            combined = (title + " " + tags).lower()
            if keyword_filter and not any(t in combined for t in keyword_filter):
                continue
            jobs.append({
                "title":           title,
                "company":         company,
                "location":        "Remote",
                "url":             j.get("url", ""),
                "source":          "remoteok",
                "remote":          True,
                "raw_description": j.get("description", "")[:2000],
            })
        return jobs
    except Exception:
        return []


def fetch_remotive_jobs(keyword_filter: list[str] | None = None) -> list[dict]:
    """Fetch remote jobs from Remotive's public API."""
    try:
        resp = requests.get(
            "https://remotive.com/api/remote-jobs",
            timeout=15,
            headers={"User-Agent": "Mozilla/5.0"},
        )
        if resp.status_code != 200:
            return []
        jobs = []
        for j in resp.json().get("jobs", []):
            title   = j.get("title", "")
            company = j.get("company_name", "")
            if not title or not company:
                continue
            tags     = " ".join(j.get("tags", [])).lower()
            combined = (title + " " + tags).lower()
            if keyword_filter and not any(t in combined for t in keyword_filter):
                continue
            desc_html = j.get("description", "")
            jobs.append({
                "title":           title,
                "company":         company,
                "location":        j.get("candidate_required_location", "Worldwide"),
                "url":             j.get("url", ""),
                "source":          "remotive",
                "remote":          True,
                "raw_description": (
                    BeautifulSoup(desc_html, "html.parser").get_text()[:2000]
                    if "<" in desc_html else desc_html[:2000]
                ),
            })
        return jobs
    except Exception:
        return []


def scrape_page_for_jobs(url: str, company_name: str) -> list[dict]:
    try:
        resp = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(resp.text, "html.parser")
        job_links = []
        for a in soup.find_all("a", href=True):
            text = a.get_text(strip=True)
            href = a["href"]
            if len(text) > 10 and any(k in href.lower() for k in ["job", "career", "position", "opening"]):
                job_links.append({"title": text, "url": href, "company": company_name, "source": "scraped"})
        return job_links[:50]
    except Exception:
        return []


# ──────────────────────────────────────────────────────────────────────────────
# Routes — Static
# ──────────────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return send_from_directory(_TEMPLATE_DIR, "index.html")


# ──────────────────────────────────────────────────────────────────────────────
# Routes — Resume
# ──────────────────────────────────────────────────────────────────────────────

@app.route("/api/resume", methods=["GET"])
def get_resume():
    text, skills = load_resume_text()
    preview = text[:500] + "..." if len(text) > 500 else text
    return jsonify({"text": preview, "skills": skills, "loaded": bool(text)})


@app.route("/api/resume/upload", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        data = request.get_json(silent=True)
        if data and "text" in data:
            text = data["text"]
        else:
            return jsonify({"error": "No file or text provided"}), 400
    else:
        f        = request.files["file"]
        filename = f.filename.lower()
        if filename.endswith(".docx"):
            from docx import Document  # type: ignore
            doc  = Document(io.BytesIO(f.read()))
            text = "\n".join(p.text for p in doc.paragraphs)
        elif filename.endswith(".txt") or filename.endswith(".md"):
            text = f.read().decode("utf-8")
        elif filename.endswith(".pdf"):
            import pdfplumber  # type: ignore
            pdf_bytes = f.read()
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            if not text.strip():
                return jsonify({"error": "Could not extract text from PDF — try a .docx or .txt instead"}), 400
        else:
            text = f.read().decode("utf-8", errors="replace")

    skills = extract_skills_from_resume(text)
    conn   = get_db()
    conn.execute(
        "INSERT OR REPLACE INTO resume_data (id, content, skills, last_updated) VALUES (1,?,?,?)",
        (text, json.dumps(skills), datetime.now().isoformat()),
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "skills": skills, "char_count": len(text)})


@app.route("/api/profile", methods=["GET"])
def get_profile():
    return jsonify(load_search_profile())


@app.route("/api/profile", methods=["PUT"])
def update_profile():
    data = request.get_json() or {}
    return jsonify({"ok": True, "profile": save_search_profile(data)})


# ──────────────────────────────────────────────────────────────────────────────
# Routes — Jobs CRUD
# ──────────────────────────────────────────────────────────────────────────────

@app.route("/api/jobs", methods=["GET"])
def get_jobs():
    status      = request.args.get("status")
    search      = request.args.get("q", "").lower()
    remote_only = request.args.get("remote") == "1"
    new_only    = request.args.get("new") == "1"

    conn   = get_db()
    query  = "SELECT * FROM jobs WHERE 1=1"
    params: list = []
    if status:
        query += " AND status=?"
        params.append(status)
    if remote_only:
        query += " AND remote=1"
    if new_only:
        query += " AND is_new=1"
    if search:
        query += " AND (LOWER(title) LIKE ? OR LOWER(company) LIKE ? OR LOWER(notes) LIKE ?)"
        params += [f"%{search}%", f"%{search}%", f"%{search}%"]
    query += " ORDER BY score DESC, date_added DESC"

    rows = conn.execute(query, params).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/jobs", methods=["POST"])
def add_job():
    data           = request.get_json()
    resume_text, resume_skills = load_resume_text()
    search_profile = load_search_profile()
    score, reasons = score_job_against_resume(
        data.get("title", ""),
        data.get("raw_description", "") + " " + data.get("notes", ""),
        resume_text, resume_skills, search_profile,
    )
    conn = get_db()
    cur  = conn.execute(
        """
        INSERT INTO jobs (title, company, location, remote, salary, url, source, status,
                          score, match_reasons, notes, date_added, date_updated, raw_description, is_new)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        """,
        (
            data.get("title", ""), data.get("company", ""), data.get("location", ""),
            1 if data.get("remote") else 0,
            data.get("salary", ""), data.get("url", ""),
            data.get("source", "manual"), data.get("status", "interested"),
            data.get("score", score), json.dumps(data.get("match_reasons", reasons)),
            data.get("notes", ""),
            datetime.now().isoformat(), datetime.now().isoformat(),
            data.get("raw_description", ""),
            1,  # manually added jobs are always "new"
        ),
    )
    job_id = cur.lastrowid
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "id": job_id, "score": score, "reasons": reasons})


@app.route("/api/jobs/<int:job_id>", methods=["PUT"])
def update_job(job_id):
    data    = request.get_json()
    conn    = get_db()
    fields  = []
    values  = []
    allowed = [
        "title", "company", "location", "remote", "salary", "url",
        "status", "score", "notes", "date_applied", "raw_description", "is_new",
    ]
    for k in allowed:
        if k in data:
            fields.append(f"{k}=?")
            values.append(data[k])
    if not fields:
        return jsonify({"error": "Nothing to update"}), 400
    fields.append("date_updated=?")
    values += [datetime.now().isoformat(), job_id]
    conn.execute(f"UPDATE jobs SET {', '.join(fields)} WHERE id=?", values)
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/jobs/<int:job_id>", methods=["DELETE"])
def delete_job(job_id):
    conn = get_db()
    conn.execute("DELETE FROM jobs WHERE id=?", (job_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/jobs/import-csv", methods=["POST"])
def import_csv_route():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file"}), 400
    content        = f.read().decode("utf-8")
    reader         = csv.DictReader(io.StringIO(content))
    resume_text, resume_skills = load_resume_text()
    search_profile = load_search_profile()
    conn           = get_db()
    imported       = 0
    for row in reader:
        title   = row.get("title", "").strip()
        company = row.get("company", "").strip()
        if not title or not company:
            continue
        if conn.execute("SELECT id FROM jobs WHERE title=? AND company=?", (title, company)).fetchone():
            continue
        csv_score = row.get("score", "")
        try:
            score   = float(csv_score)
            reasons = [r.strip() for r in row.get("reasons", "").split("|") if r.strip()]
        except (ValueError, TypeError):
            score, reasons = score_job_against_resume(title, "", resume_text, resume_skills, search_profile)
        location   = row.get("location", "")
        remote_val = str(row.get("remote", "")).lower() in ("true", "1", "yes", "remote")
        conn.execute(
            """
            INSERT INTO jobs (title, company, location, remote, url, source, status,
                              score, match_reasons, date_added, date_updated, is_new)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                title, company, location, 1 if remote_val else 0,
                row.get("url", ""), row.get("source", "csv"), "interested",
                score, json.dumps(reasons),
                datetime.now().isoformat(), datetime.now().isoformat(),
                1,
            ),
        )
        imported += 1
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "imported": imported})


# ──────────────────────────────────────────────────────────────────────────────
# Routes — Export
# ──────────────────────────────────────────────────────────────────────────────

@app.route("/api/jobs/export", methods=["GET"])
def export_jobs_csv():
    conn = get_db()
    rows = conn.execute(
        """
        SELECT title, company, location, remote, salary, url, source, status,
               score, match_reasons, notes, date_added, date_applied, date_updated, is_new
        FROM jobs
        ORDER BY score DESC, date_added DESC
        """
    ).fetchall()
    conn.close()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "Title", "Company", "Location", "Remote", "Salary", "URL", "Source",
        "Status", "Match Score", "Match Reasons", "Competition Level", "Notes",
        "Date Added", "Date Applied", "Last Updated", "New",
    ])
    for r in rows:
        try:
            reasons = ", ".join(json.loads(r["match_reasons"] or "[]"))
        except Exception:
            reasons = r["match_reasons"] or ""
        writer.writerow([
            r["title"], r["company"], r["location"] or "",
            "Yes" if r["remote"] else "No",
            r["salary"] or "", r["url"] or "", r["source"] or "",
            r["status"] or "", r["score"] or "",
            reasons,
            estimate_competition(r["title"], r["company"]),
            r["notes"] or "",
            (r["date_added"] or "")[:10],
            (r["date_applied"] or "")[:10],
            (r["date_updated"] or "")[:10],
            "Yes" if r["is_new"] else "No",
        ])
    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=jobtracker_export.csv"},
    )


# ──────────────────────────────────────────────────────────────────────────────
# Routes — Sources & Fetch
# ──────────────────────────────────────────────────────────────────────────────

@app.route("/api/sources", methods=["GET"])
def get_sources():
    conn = get_db()
    rows = conn.execute("SELECT * FROM sources ORDER BY active DESC, name").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/sources", methods=["POST"])
def add_source():
    data = request.get_json()
    conn = get_db()
    cur  = conn.execute(
        "INSERT INTO sources (name, url, type, active) VALUES (?,?,?,1)",
        (data.get("name", ""), data.get("url", ""), data.get("type", "direct")),
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "id": cur.lastrowid})


@app.route("/api/sources/<int:src_id>", methods=["PUT"])
def update_source(src_id):
    data = request.get_json()
    conn = get_db()
    conn.execute(
        "UPDATE sources SET name=?, url=?, type=?, active=? WHERE id=?",
        (data["name"], data["url"], data.get("type", "direct"), data.get("active", 1), src_id),
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/sources/<int:src_id>", methods=["DELETE"])
def delete_source(src_id):
    conn = get_db()
    conn.execute("DELETE FROM sources WHERE id=?", (src_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/sources/<int:src_id>/fetch", methods=["POST"])
def fetch_source(src_id):
    conn   = get_db()
    source = conn.execute("SELECT * FROM sources WHERE id=?", (src_id,)).fetchone()
    if not source:
        conn.close()
        return jsonify({"error": "Source not found"}), 404

    source         = dict(source)
    resume_text, resume_skills = load_resume_text()
    search_profile = load_search_profile()
    jobs: list[dict] = []

    stype = source["type"]
    if stype == "greenhouse":
        jobs = fetch_greenhouse_jobs(source["name"], source["url"])
    elif stype == "lever":
        jobs = fetch_lever_jobs(source["name"], source["url"])
    elif stype == "ashby":
        jobs = fetch_ashby_jobs(source["name"], source["url"])
    elif stype == "remoteok":
        jobs = fetch_remoteok_jobs()
    elif stype == "remotive":
        jobs = fetch_remotive_jobs()
    else:
        jobs = scrape_page_for_jobs(source["url"], source["name"])

    added = 0
    for j in jobs:
        if conn.execute(
            "SELECT id FROM jobs WHERE title=? AND company=?",
            (j["title"], j["company"]),
        ).fetchone():
            continue

        score, reasons = score_job_against_resume(
            j["title"], j.get("raw_description", ""), resume_text, resume_skills, search_profile
        )
        if is_filtered_match(score, reasons):
            continue

        location   = j.get("location", "")
        remote_val = any(k in location.lower() for k in ["remote", "anywhere"]) or j.get("remote", False)

        conn.execute(
            """
            INSERT INTO jobs (title, company, location, remote, url, source, status,
                              score, match_reasons, date_added, date_updated, raw_description, is_new)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                j["title"], j["company"], location, 1 if remote_val else 0,
                j.get("url", ""), j.get("source", stype), "interested",
                score, json.dumps(reasons),
                datetime.now().isoformat(), datetime.now().isoformat(),
                j.get("raw_description", "")[:3000],
                1,  # mark as new — seen for the first time this run
            ),
        )
        added += 1

    conn.execute(
        "UPDATE sources SET last_fetched=?, job_count=? WHERE id=?",
        (datetime.now().isoformat(), len(jobs), src_id),
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "fetched": len(jobs), "added": added})


# ──────────────────────────────────────────────────────────────────────────────
# Routes — Rescore / Stats
# ──────────────────────────────────────────────────────────────────────────────

@app.route("/api/jobs/rescore", methods=["POST"])
def rescore_all_jobs():
    resume_text, resume_skills = load_resume_text()
    if not resume_text:
        return jsonify({"error": "No resume loaded — upload your resume first"}), 400
    search_profile = load_search_profile()
    conn           = get_db()
    rows           = conn.execute("SELECT id, title, raw_description, notes FROM jobs").fetchall()
    updated = zeroed = 0
    for r in rows:
        score, reasons = score_job_against_resume(
            r["title"],
            (r["raw_description"] or "") + " " + (r["notes"] or ""),
            resume_text, resume_skills, search_profile,
        )
        conn.execute(
            "UPDATE jobs SET score=?, match_reasons=?, date_updated=? WHERE id=?",
            (score, json.dumps(reasons), datetime.now().isoformat(), r["id"]),
        )
        updated += 1
        if score == 0:
            zeroed += 1
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "rescored": updated, "zeroed_out": zeroed})


@app.route("/api/stats", methods=["GET"])
def get_stats():
    conn       = get_db()
    total      = conn.execute("SELECT COUNT(*) FROM jobs").fetchone()[0]
    new_count  = conn.execute("SELECT COUNT(*) FROM jobs WHERE is_new=1").fetchone()[0]
    by_status  = conn.execute("SELECT status, COUNT(*) as n FROM jobs GROUP BY status").fetchall()
    top_co     = conn.execute(
        "SELECT company, COUNT(*) as n FROM jobs GROUP BY company ORDER BY n DESC LIMIT 5"
    ).fetchall()
    conn.close()
    return jsonify({
        "total":         total,
        "new_count":     new_count,
        "by_status":     {r["status"]: r["n"] for r in by_status},
        "top_companies": [{"company": r["company"], "count": r["n"]} for r in top_co],
    })


# ──────────────────────────────────────────────────────────────────────────────
# Routes — URL scrape / ATS detect / Smart fetch
# ──────────────────────────────────────────────────────────────────────────────

@app.route("/api/scrape-url", methods=["POST"])
def scrape_url():
    data = request.get_json()
    url  = data.get("url", "")
    if not url:
        return jsonify({"error": "No URL"}), 400
    try:
        resp = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        text  = soup.get_text(separator="\n", strip=True)[:5000]
        title = ""
        for tag in ["h1", "h2"]:
            el = soup.find(tag)
            if el:
                title = el.get_text(strip=True)
                break
        return jsonify({"ok": True, "title": title, "description": text})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/detect-ats", methods=["POST"])
def detect_ats():
    data = request.get_json()
    name = data.get("name", "").strip()
    url  = data.get("url", "").strip()
    slug = name.lower().replace(" ", "").replace("-", "").replace(".", "")
    for pattern, ats in [
        (r"greenhouse\.io/([^/?#]+)",  "greenhouse"),
        (r"lever\.co/([^/?#]+)",       "lever"),
        (r"ashbyhq\.com/([^/?#]+)",    "ashby"),
        (r"remoteok\.com",             "remoteok"),
    ]:
        m = re.search(pattern, url)
        if m:
            return jsonify({"ats": ats, "slug": m.group(1) if m.lastindex else slug, "url": url})
    probes = [
        ("greenhouse", f"https://boards-api.greenhouse.io/v1/boards/{slug}/jobs"),
        ("lever",      f"https://api.lever.co/v0/postings/{slug}?mode=json"),
        ("ashby",      f"https://api.ashbyhq.com/posting-api/job-board/{slug}"),
    ]
    for ats, probe_url in probes:
        try:
            r = requests.head(probe_url, timeout=6, headers={"User-Agent": "Mozilla/5.0"})
            if r.status_code == 200:
                return jsonify({"ats": ats, "slug": slug, "url": probe_url})
        except Exception:
            continue
    return jsonify({"ats": "unknown", "slug": slug})


@app.route("/api/smart-fetch", methods=["POST"])
def smart_fetch():
    from urllib.parse import urlparse

    data = request.get_json()
    raw  = (data.get("input", "") or "").strip()
    if not raw:
        return jsonify({"error": "No input provided"}), 400

    parsed     = urlparse(raw if "://" in raw else "https://" + raw)
    hostname   = parsed.hostname or ""
    path_parts = [p for p in parsed.path.strip("/").split("/") if p]

    def clean(s: str) -> str:
        return re.sub(r"[^a-z0-9]", "", s.lower())

    slug_candidates: list[str] = []
    known_ats = {
        "boards.greenhouse.io": "greenhouse",
        "greenhouse.io":        "greenhouse",
        "jobs.lever.co":        "lever",
        "lever.co":             "lever",
        "jobs.ashbyhq.com":     "ashby",
        "ashbyhq.com":          "ashby",
    }
    detected_ats = None
    for host_pat, ats_name in known_ats.items():
        if host_pat in hostname:
            detected_ats = ats_name
            if path_parts:
                slug_candidates.insert(0, clean(path_parts[0]))
            break

    domain_base = re.sub(r"^(www|jobs|careers|app)\.", "", hostname)
    domain_base = re.sub(r"\.(com|io|co|net|org|ai|app)(\..*)?$", "", domain_base)
    slug_candidates.append(clean(domain_base))
    slug_candidates.append(clean(re.sub(r"https?://", "", raw).split("/")[0].split(".")[0]))

    seen: set[str] = set()
    slugs: list[str] = []
    for s in slug_candidates:
        if s and s not in seen:
            seen.add(s)
            slugs.append(s)

    company_name = data.get("name") or domain_base.replace("-", " ").title() or (slugs[0].title() if slugs else raw)

    ats_probes = [
        ("greenhouse", lambda sl: f"https://boards-api.greenhouse.io/v1/boards/{sl}/jobs"),
        ("lever",      lambda sl: f"https://api.lever.co/v0/postings/{sl}?mode=json"),
        ("ashby",      lambda sl: f"https://api.ashbyhq.com/posting-api/job-board/{sl}"),
    ]
    found_ats  = detected_ats
    found_slug = slugs[0] if slugs else ""
    found_url  = raw

    if not found_ats:
        for ats_name, url_fn in ats_probes:
            for sl in slugs:
                try:
                    r = requests.head(url_fn(sl), timeout=6, headers={"User-Agent": "Mozilla/5.0"})
                    if r.status_code == 200:
                        found_ats  = ats_name
                        found_slug = sl
                        found_url  = url_fn(sl)
                        break
                except Exception:
                    continue
            if found_ats:
                break

    jobs_raw: list[dict] = []
    if found_ats == "greenhouse":
        jobs_raw = fetch_greenhouse_jobs(company_name, found_url)
    elif found_ats == "lever":
        jobs_raw = fetch_lever_jobs(company_name, found_url)
    elif found_ats == "ashby":
        jobs_raw = fetch_ashby_jobs(company_name, found_url)
    else:
        career_paths = ["/careers", "/jobs", "/join", "/work-with-us", "/about/careers"]
        base = f"https://{hostname}" if hostname else raw
        for cp in career_paths:
            try:
                r = requests.get(base + cp, timeout=8,
                                 headers={"User-Agent": "Mozilla/5.0"}, allow_redirects=True)
                if r.status_code == 200:
                    found_url  = r.url
                    jobs_raw   = scrape_page_for_jobs(found_url, company_name)
                    found_ats  = "direct"
                    break
            except Exception:
                continue

    resume_text, resume_skills = load_resume_text()
    search_profile = load_search_profile()
    conn = get_db()

    src_row = conn.execute("SELECT id FROM sources WHERE name=?", (company_name,)).fetchone()
    if src_row:
        src_id = src_row["id"]
    else:
        cur = conn.execute(
            "INSERT INTO sources (name, url, type, active) VALUES (?,?,?,1)",
            (company_name, found_url, found_ats or "direct"),
        )
        src_id = cur.lastrowid

    added   = 0
    results = []
    for j in jobs_raw:
        score, reasons = score_job_against_resume(
            j["title"], j.get("raw_description", ""), resume_text, resume_skills, search_profile
        )
        if is_filtered_match(score, reasons):
            continue
        location   = j.get("location", "")
        remote_val = any(k in location.lower() for k in ["remote", "anywhere"]) or j.get("remote", False)

        if not conn.execute(
            "SELECT id FROM jobs WHERE title=? AND company=?", (j["title"], j["company"])
        ).fetchone():
            conn.execute(
                """
                INSERT INTO jobs (title, company, location, remote, url, source, status,
                                  score, match_reasons, date_added, date_updated, raw_description, is_new)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    j["title"], j["company"], location, 1 if remote_val else 0,
                    j.get("url", ""), found_ats or "direct", "interested",
                    score, json.dumps(reasons),
                    datetime.now().isoformat(), datetime.now().isoformat(),
                    j.get("raw_description", "")[:3000],
                    1,
                ),
            )
            added += 1

        results.append({
            "title":   j["title"],
            "company": j["company"],
            "score":   score,
            "reasons": reasons,
            "url":     j.get("url", ""),
        })

    conn.execute("UPDATE sources SET last_fetched=?, job_count=? WHERE id=?",
                 (datetime.now().isoformat(), len(jobs_raw), src_id))
    conn.commit()
    conn.close()

    results.sort(key=lambda x: x["score"], reverse=True)
    return jsonify({
        "ok":          True,
        "ats":         found_ats or "unknown",
        "company":     company_name,
        "source_id":   src_id,
        "fetched":     len(jobs_raw),
        "added":       added,
        "top_matches": results[:10],
    })


@app.route("/api/jobboard-search-urls", methods=["GET"])
def jobboard_search_urls():
    q = request.args.get("q", "").strip()
    if not q:
        _, resume_skills = load_resume_text()
        abbrev_skills = [s.upper() for s in resume_skills if s.upper() == s and len(s) <= 5]
        word_skills   = [s for s in resume_skills if not (s.upper() == s and len(s) <= 5)]
        top_terms     = (abbrev_skills[:3] + word_skills[:2]) or ["jobs"]
        q = " ".join(top_terms[:4])

    eq      = requests.utils.quote(q)
    eq_plus = q.replace(" ", "+")

    boards = [
        {"name": "LinkedIn Jobs",     "icon": "in", "color": "#0a66c2",
         "url": f"https://www.linkedin.com/jobs/search/?keywords={eq_plus}",
         "note": "Add Remote filter in LinkedIn UI"},
        {"name": "Indeed",            "icon": "in", "color": "#2164f3",
         "url": f"https://www.indeed.com/jobs?q={eq_plus}&l=Remote",
         "note": "Remote filtered"},
        {"name": "Glassdoor",         "icon": "gd", "color": "#0caa41",
         "url": f"https://www.glassdoor.com/Job/jobs.htm?sc.keyword={eq}&locT=N&remoteWorkType=1",
         "note": "Remote filtered"},
        {"name": "Wellfound",         "icon": "wf", "color": "#000",
         "url": f"https://wellfound.com/jobs?q={eq_plus}&remote=true",
         "note": "Startup-focused"},
        {"name": "Built In",          "icon": "bi", "color": "#5c36f3",
         "url": f"https://builtin.com/jobs?search={eq_plus}&remote=true",
         "note": "Tech company focused"},
        {"name": "ZipRecruiter",      "icon": "zr", "color": "#4a90d9",
         "url": f"https://www.ziprecruiter.com/jobs-search?search={eq_plus}&location=Remote",
         "note": "Remote filtered"},
        {"name": "FlexJobs",          "icon": "fj", "color": "#00b373",
         "url": f"https://www.flexjobs.com/search?search={eq_plus}&location=Remote",
         "note": "Vetted remote-only"},
        {"name": "RemoteOK",          "icon": "ro", "color": "#00d4aa",
         "url": f"https://remoteok.com/remote-{eq_plus.replace('+', '-')}-jobs",
         "note": "Remote-only board"},
        {"name": "USAJobs",           "icon": "us", "color": "#1a3c6e",
         "url": f"https://www.usajobs.gov/search/results/?k={eq_plus}&p=1",
         "note": "Federal / government roles"},
    ]
    return jsonify(boards)


# ──────────────────────────────────────────────────────────────────────────────
# Entry point
# ──────────────────────────────────────────────────────────────────────────────

def _auto_load_resume() -> None:
    """Try to auto-import a resume from the data/ folder on first launch."""
    resume_text, _ = load_resume_text()
    if resume_text:
        return
    candidates = [
        ("data/resume.pdf",  "pdf"),
        ("data/resume.docx", "docx"),
        ("data/resume.txt",  "txt"),
        ("data/resume.md",   "txt"),
    ]
    for fpath, ftype in candidates:
        if not os.path.exists(fpath):
            continue
        try:
            if ftype == "docx":
                from docx import Document  # type: ignore
                doc  = Document(fpath)
                text = "\n".join(p.text for p in doc.paragraphs)
            elif ftype == "pdf":
                import pdfplumber  # type: ignore
                with pdfplumber.open(fpath) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            else:
                with open(fpath, encoding="utf-8") as fh:
                    text = fh.read()
            if text.strip():
                skills = extract_skills_from_resume(text)
                conn   = get_db()
                conn.execute(
                    "INSERT OR REPLACE INTO resume_data (id, content, skills, last_updated) VALUES (1,?,?,?)",
                    (text, json.dumps(skills), datetime.now().isoformat()),
                )
                conn.commit()
                conn.close()
                print(f"✓ Resume auto-loaded from {fpath} ({len(text)} chars, {len(skills)} skills)")
                return
        except Exception as e:
            print(f"Resume auto-load ({fpath}) skipped: {e}")
    print("ℹ  No resume found in data/ — upload via the UI to enable job scoring")


def main() -> None:
    init_db()
    _auto_load_resume()
    print("\n🚀  Job Agent web UI running at http://localhost:5001\n")
    app.run(debug=False, host="0.0.0.0", port=5001)


if __name__ == "__main__":
    main()
